import streamlit as st
import os
import PyPDF2
import openai
import io
import pandas as pd
from docx import Document
from dotenv import load_dotenv
from utilities.chroma import create_chroma_collection, store_chunks_in_chroma_collection
from utilities.chunk_by_sentences import chunk_by_sentences
from utilities.llm_agent import agent_with_search

load_dotenv()

if "collection" not in st.session_state:
    st.session_state.collection = None

if "chunks" not in st.session_state:
    st.session_state.chunks = None

if "report" not in st.session_state:
    st.session_state.report = None

if "clean_csv" not in st.session_state:
    st.session_state.clean_csv = None

# Title
st.title("Ask Hukumi 2.0")

# Upload documents
st.subheader("Upload Documents")
uploaded_documents = st.file_uploader("Choose PDF file(s) to upload", type=["pdf"], accept_multiple_files=True)

if not uploaded_documents:
    st.session_state.collection = None
    st.session_state.chunks = None
    st.session_state.report = None
    st.session_state.clean_csv = None

# Create Chroma collection
collection_name = st.text_input("What is the name of your collection?")

if uploaded_documents and collection_name and st.button("Upload and Create Collection"):
    collection = create_chroma_collection(f"{collection_name}")
    st.session_state.collection = collection

    # Store documents
    with st.spinner("Uploading and processing document(s)..."):

        document_folder = "document_folder"

        if not os.path.exists(document_folder):
            os.makedirs(document_folder)

        for uploaded_document in uploaded_documents:
            file_name = uploaded_document.name
            file_path = os.path.join(document_folder, file_name)
            
            with open(file_path, "wb") as f:
                f.write(uploaded_document.getbuffer())
                
            st.success(f"PDF file '{file_name}' has been uploaded successfully")

            # Read documents
            try:
                with open(file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    num_pages = len(pdf_reader.pages)
                    text = ""
            
                    for page in range(num_pages):
                        text += pdf_reader.pages[page].extract_text() + "\n\n"
        
                # Chunk documents
                chunks = chunk_by_sentences(text)
                
                # Store chunks in Chroma
                for idx, chunk in enumerate(chunks):
                    doc_id = f"{file_name}: {idx}"
                    #st.write(f"Trying to add doc_id: {doc_id}")

                    store_chunks_in_chroma_collection(
                        collection=st.session_state.collection,
                        document_name=file_name,
                        document_chunk=chunk['content'],
                        document_id=doc_id
                    )
                
                st.info(f"Created and stored {len(chunks)} chunks from '{file_name}'")

                # Show chunks
                with st.expander(f"Document Chunks from '{file_name}' ({len(chunks)} chunks)"):
                    for i, chunk in enumerate(chunks, 1):
                        with st.expander(chunk['header']):
                            st.write(chunk['content'])
                
                st.session_state.chunks = chunks

            except Exception as e:
                st.error(f"Error processing '{file_name}': {e}")

# User objective
st.subheader("User Objective")

user_sector = st.text_input("What is your sector?")
user_company_type = st.text_input("What is your company type?")
user_jurisdiction = st.text_input("What is your jurisdiction?")
user_specific_objective = st.text_input("""What is your specific objective?
Example: 'To identify key compliance checklist for a fintech startup planning to enter Australian market'""")
user_notes = st.text_area("Other notes for the agent? If there is not any, enter 'None'")

# Number of loops
#user_number_of_loops = st.number_input("How many search iterations should the agent perform?", min_value=1, max_value=10, value=3)
user_number_of_loops = st.selectbox(
    "How many search iterations should the agent perform?",
    ("Let the agent decide", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10"),
    index=None
)

# Set user objective
if st.session_state.chunks and user_number_of_loops and st.button("Set Objective and Run Agent"):
    
    if not user_sector or not user_company_type or not user_jurisdiction or not user_specific_objective or not user_notes:
        st.warning("Please fill in all fields to set user objective")

    else:
        user_objective = f"""Perform a detailed analysis of the document(s) to identify and extract all relevant compliance checklist and obligations for a company.
        If there is more than one document, you must synthesise information across all documents to provide a comprehensive analysis.
        Do not make up any information; only use what is contained in the document(s).
        
        The context for your analysis is as follows:
        1. The company is operating in the {user_sector} sector(s).
        2. The company's type is {user_company_type}.
        3. The company operates in the jurisdiction of {user_jurisdiction}.
        4. The user's specific objective is: {user_specific_objective}.
        5. Additional notes for the agent are: {user_notes}.
        
        You must consider the following tasks when performing your analysis:
        1. Search for obligations that are specifically relevant to the company's sector(s), type, and jurisdiction.
        Follow the user's specific objective and notes.
        2. Search for reporting or periodic obligations.
        3. Search for licensing requirements or licences that the company may obtain.
        4. Search for prohibitions that the company must comply with."""
        
        st.success("User objective has been set successfully")
        st.info("Running agent... This may take a few minutes")

        # Run agent
        try:
            agent_still_searching = True
            input_list = [{"role": "user", "content": user_objective}]
            number_of_loops = 0

            if user_number_of_loops == "Let the agent decide":
                while agent_still_searching:

                    status_text = st.empty()
                    status_text.text(f"Search iteration {number_of_loops + 1}")
                    
                    number_of_loops += 1

                    input_list = agent_with_search(st.session_state.collection, user_objective, input_list)
                
                    st.write(input_list)

                    if input_list[-1].get("content")[0].get("text") == "FINISHED":
                        agent_still_searching = False

            else:
                max_loops = int(user_number_of_loops)

                # Progress setup
                progress_bar = st.progress(0)
                status_text = st.empty()

                while agent_still_searching and number_of_loops < max_loops:

                    progress = (number_of_loops + 1) / max_loops
                    progress_bar.progress(progress)
                    status_text.text(f"Search iteration {number_of_loops + 1}/{max_loops}")
                    
                    number_of_loops += 1
                    
                    input_list = agent_with_search(st.session_state.collection, user_objective, input_list)

                    st.write(input_list)

                    if number_of_loops == max_loops:
                        agent_still_searching = False

            if not agent_still_searching:

                # Generate report
                def generate_report():
                    try:
                        with st.spinner("Generating compliance report..."):
                            response = openai.responses.create(
                                model="gpt-5",
                                #response_format={"type": "text"},
                                instructions=(
                                    f"""Your objective is: {user_objective}.
                                    You have retrieved document chunks from ChromaDB searches.
                                    Identify and extract all compliance checklist and obligations relevant to the objective and make a report.
                                    For your report, you must summarise, cite, and reference the actual content of these document chunks.
                                    Include direct quotes or paraphrased information from the documents where relevant.
                                    Make it clear which information comes from which document, part, section, etc.
                                    Again, you must synthesise information across all documents to provide a comprehensive report.
                                    Your final response must be a report that the user can read and understand easily.
                                    Do not make up any information; only use what is contained in the document chunks provided to you.
                                    Do not ask any questions or offer anything to the user and say 'End of report.' at the end of the response."""
                                ),
                                input=input_list
                            )

                            # Extract text safely even if output_text is None
                            report = getattr(response, "output_text", None)

                            # Fallback: try extracting manually if output_text is missing
                            if not report and hasattr(response, "output"):
                                text_outputs = [
                                    item.get("text")
                                    for item in response.output
                                    if isinstance(item, dict) and item.get("type") == "output_text"
                                ]
                                report = "\n".join(text_outputs) if text_outputs else None

                            # Final fallback
                            if not report or len(report.strip()) == len("End of report."):
                                report = "No report was generated. The model may have produced empty or function-only output"
                            
                            return report

                    except Exception as e:
                        st.error(f"Error generating report: {e}")

                report = generate_report()
                st.session_state.report = report

                if not st.session_state.report or st.session_state.report == "No report was generated. The model may have produced empty or function-only output":
                    st.info("You can retry generating the report")
                    if st.button("Retry Generating Report"):
                        report = generate_report()
                        st.session_state.report = report

        except Exception as e:
            st.error(f"Error running agent: {e}")

# Show report
if st.session_state.report:
    st.subheader("Compliance Report")
    st.write(st.session_state.report)

    # Download report
    report_text = str(st.session_state.report)

    doc = Document()
    doc.add_heading('Compliance Report', 0)
    doc.add_paragraph(report_text)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    st.download_button(
        label="Download Report",
        data=doc_io.getvalue(),
        file_name=f"report_{collection_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Generate CSV
    if st.button("Generate CSV"):
        with st.spinner("Generating compliance checklist CSV..."):
            csv_text = openai.responses.create(
                model="gpt-5",
                input=[{
                    "role": "user",
                    "content": f"""Based on the report provided below, create compliance checklist in CSV format.
                    The checklist must include the following columns:
                    'Obligation', 'Description', 'Type', 'Source Document', 'Section/Part', 'Deadline/Frequency', 'Responsible Party', and 'Notes'.
                    'Type' can include 'Continuous', 'Reporting', 'Licensing', and 'Prohibition'.
                    Most obligations are 'Continuous', except if it is 'Reporting', 'Licensing', or 'Prohibition'.
                    Ensure that each obligation is clearly defined and includes all relevant details from the report.
                    Provide the output strictly in CSV format without any additional text or explanation.

                    Report:
                    {st.session_state.report}"""
                }]
            )
            
            # Fix CSV function
            def fix_csv(messy_csv):
           
                # Fix encoding to UTF-8
                if isinstance(messy_csv, bytes):
                    messy_csv = messy_csv.decode('utf-8', errors='ignore')
                
                # Remove markdown
                if "```" in messy_csv:
                    parts = messy_csv.split("```")
                    for part in parts:
                        if "," in part and "\n" in part:
                            messy_csv = part.replace("csv", "").strip()
                            break
                
                # Fix common encoding issues
                messy_csv = messy_csv.encode('utf-8', errors='ignore').decode('utf-8')
                
                # Clean up
                clean_csv = messy_csv.strip()
                
                return clean_csv

            # Use the function
            csv_data = csv_text.output_text
            clean_csv = fix_csv(csv_data)
            st.session_state.clean_csv = clean_csv

# Preview and download CSV
if st.session_state.clean_csv:
    st.subheader("Compliance Checklist CSV")

    try:
        df = pd.read_csv(io.StringIO(st.session_state.clean_csv))
        st.dataframe(df)
        
        st.download_button(
            label="Download CSV",
            data=st.session_state.clean_csv.encode('utf-8'),
            file_name=f"csv_{collection_name}.csv",
            mime="text/csv"
        )

    # If still broken, download as-is
    except:
        st.warning("CSV might need manual cleanup")
        st.download_button(
            label="Download Raw CSV",
            data=st.session_state.clean_csv.encode('utf-8'),
            file_name=f"raw_csv_{collection_name}.csv",
            mime="text/csv"
        )